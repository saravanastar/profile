
# Refactored: Read profile data from JSON (ask.json or poorni.json)
import os
import json
import sys
import shutil
from docx import Document
import subprocess
import platform

def convert_to_pdf(docx_path):
    pdf_path = docx_path.replace('.docx', '.pdf')
    output_dir = os.path.dirname(docx_path) or '.'
    system = platform.system().lower()
    cmd = None
    
    if system == 'darwin':  # macOS
        converter = shutil.which('soffice') or shutil.which('libreoffice')
        if converter is None:
            print("Warning: LibreOffice/soffice not found. Skipping PDF conversion and keeping the DOCX output.")
            return
        cmd = [converter, '--headless', '--convert-to', 'pdf', docx_path, '--outdir', output_dir]
    elif system == 'linux':
        converter = shutil.which('libreoffice') or shutil.which('soffice')
        if converter is None:
            print("Warning: LibreOffice/soffice not found. Skipping PDF conversion and keeping the DOCX output.")
            return
        cmd = [converter, '--headless', '--convert-to', 'pdf', docx_path, '--outdir', output_dir]
    elif system == 'windows':
        try:
            from docx2pdf import convert
            convert(docx_path)
            return
        except Exception as e:
            print(f"Warning: docx2pdf failed: {e}")
            converter = shutil.which('soffice') or shutil.which('libreoffice')
            if converter is None:
                print("Warning: LibreOffice/soffice not found. Skipping PDF conversion and keeping the DOCX output.")
                return
            cmd = [converter, '--headless', '--convert-to', 'pdf', docx_path, '--outdir', output_dir]
    
    if cmd is None:
        print(f"Warning: Unsupported OS '{system}'. Skipping PDF conversion; DOCX kept.")
        return
    try:
        subprocess.run(cmd, check=True)
        print(f"PDF created successfully: {pdf_path}")
    except subprocess.CalledProcessError as e:
        print(f"Error converting to PDF: {e}")
        print("Please ensure LibreOffice is installed:")
        if system == 'darwin':
            print("brew install --cask libreoffice")
        elif system == 'linux':
            print("sudo apt-get install libreoffice")
        elif system == 'windows':
            print("Download from https://www.libreoffice.org/download/download/")

def load_profile_data(profile_json=None):
    if profile_json is None:
        profile_json = os.environ.get("PROFILE_JSON")
        if profile_json is None:
            raise ValueError("Profile JSON file name must be provided as a parameter or via the PROFILE_JSON environment variable.")
    data_dir = os.path.join(os.path.dirname(__file__), "data")
    json_path = os.path.join(data_dir, profile_json)
    if not os.path.exists(json_path):
        raise FileNotFoundError(f"Profile JSON file not found: {json_path}")
    with open(json_path, "r") as f:
        profile = json.load(f)
    return profile

profile_json = sys.argv[1] if len(sys.argv) > 1 else None
# Usage: pass the file name as a parameter or set PROFILE_JSON env variable
profile = load_profile_data(profile_json)
repo_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
doc = Document()

# Name and Contact

doc.add_heading(profile["name"], level=1)
contact = profile.get("contact", {})
city = contact.get("city", "")
state = contact.get("state", "")
mobile = contact.get("mobileNumber", "")
email = contact.get("email", "")
location = ", ".join(part for part in [city, state] if part)
contact_parts = [p for p in [location, mobile, email] if p]
doc.add_paragraph(" | ".join(contact_parts))
# Social links
social_lines = []
for s in profile.get("social", []):
    for k, v in s.items():
        social_lines.append(f"{v.get('name')}: {(v.get('url') or '').strip()}")
for o in profile.get("other", []):
    for k, v in o.items():
        social_lines.append(f"{v.get('name')}: {(v.get('url') or '').strip()}")
for line in social_lines:
    doc.add_paragraph(line)
# Summary
summary = profile.get("profileSummary", {})
doc.add_heading(summary.get("title", "Professional Summary"), level=2)
doc.add_paragraph(summary.get("content", ""))
# Skills
doc.add_heading("Skills", level=2)
for skill in profile.get("skills", []):
    para = doc.add_paragraph()
    run_title = para.add_run(f"{skill['title']}:")
    run_title.bold = True
    para.add_run(f" {skill['text']}")
# Experience
doc.add_heading("Professional Experience", level=2)
for exp in profile.get("experiences", []):
    doc.add_heading(f"{exp['title']} – {exp['company']}", level=3)
    doc.add_paragraph(f"{exp['location']} | {exp['dates']}")
    for bullet in exp.get("bullets", []):
        doc.add_paragraph(bullet, style="List Bullet")
# Projects (optional, data-driven)
projects = profile.get("projects", [])
if projects:
    doc.add_heading("Projects", level=2)
    for proj in projects:
        para = doc.add_paragraph()
        name_run = para.add_run(proj.get("name", ""))
        name_run.bold = True
        dates = proj.get("dates", "")
        if dates:
            para.add_run(f" ({dates})")
        desc = proj.get("description", "")
        if desc:
            para.add_run(f" \u2014 {desc}")
        link = (proj.get("link") or "").strip()
        if link:
            doc.add_paragraph(link)

# Education & Certifications (data-driven with fallbacks)
doc.add_heading("Education & Certifications", level=2)
education = profile.get("education", [])
if education:
    for edu in education:
        left = " \u2013 ".join(part for part in [edu.get("degree", ""), edu.get("field", "")] if part)
        segments = [seg for seg in [left, edu.get("institution", ""), edu.get("dates", "")] if seg]
        doc.add_paragraph(", ".join(segments))
else:
    doc.add_paragraph("Bachelor of Engineering \u2013 Computer Science, Anna University")

certifications = profile.get("certifications", [])
if certifications:
    for cert in certifications:
        doc.add_paragraph(cert, style="List Bullet")
else:
    doc.add_paragraph("AWS Certified Solutions Architect \u2013 Associate")

file_path = os.path.join(repo_root, f"{profile['name'].replace(' ', '_')}_Resume.docx")


doc.save(file_path)
print("Converting to PDF...")
convert_to_pdf(file_path)