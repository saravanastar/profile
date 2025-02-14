# Correcting the formatting and retrying document creation
from docx import Document
from docx2pdf import convert

# Create a new Document
doc = Document()

# Add Name and Contact Information
doc.add_heading("Saravanakumar Arunachalam", level=1)
doc.add_paragraph("Charlotte, NC | +1 214-713-5765 | saravanastar@gmail.com")
doc.add_paragraph("LinkedIn: https://www.linkedin.com/in/saravanakumar-arunachalam-a0173023")
doc.add_paragraph("GitHub: https://github.com/saravanastar | Medium: https://medium.com/@saravanastar")

# Add Professional Summary
doc.add_heading("Professional Summary", level=2)
doc.add_paragraph(
    "Software engineering professional with over 14 years of experience specializing in cloud-native, "
    "microservice, and event-driven architectures. Proficient in Java, Go, Python, Spring, Kubernetes, and AWS, "
    "with expertise in developing scalable, secure, and high-performance solutions. Strong background in DevOps, "
    "CI/CD, and infrastructure automation using Terraform and Helm. Passionate about driving innovation and solving "
    "complex challenges in high-volume, high-velocity data systems."
)

# Add Skills Section
doc.add_heading("Skills", level=2)
skills = [
    {
        "title": "Backend Development",
        "text": "Java, Go, Python, NodeJS, Spring, Microservices, Event-Driven Design"
    },
    {
        "title": "Cloud & DevOps",
        "text": "AWS, Kubernetes, Docker, Terraform, Helm, CloudFormation, AWS CDK"
    },
    {
        "title": "Streaming & Messaging",
        "text": "Kafka, Kinesis, SQS, SNS, Flink"
    },
    {
        "title": "Frontend Development",
        "text": "HTML, CSS, JavaScript, Angular, ReactJS"
    },
    {
        "title": "Databases & Storage",
        "text": "Oracle, Postgres, MySQL, MongoDB, ElasticSearch, DynamoDB, S3"
    },
    {
        "title": "Design & Architecture",
        "text": "DDD, CQRS, Circuit Breaker, Saga, Distributed Systems"
    },
    {
        "title": "Development Practices",
        "text": "SOLID, TDD, Pair Programming, Agile (Scrum/Kanban)"
    }
]
for skill in skills:
    para = doc.add_paragraph()
    run_title = para.add_run(f"{skill['title']}:")
    run_title.bold = True
    para.add_run(f" {skill['text']}")

# Add Professional Experience Section
doc.add_heading("Professional Experience", level=2)

# Define roles with detailed entries
roles = [
    {
        "title": "Senior Engineer - Platform",
        "company": "Geico",
        "location": "Charlotte, NC",
        "dates": "Jun 2024 – Present",
        "bullets": [
            "Designed and developed a serverless platform leveraging Knative, KEDA, Istio, and cert-manager on "
            "Kubernetes.",
            "Deployed container-as-a-service (CaaS) and function-as-a-service (FaaS) solutions to support scalable, "
            "performance-driven architectures.",
            "Led cross-functional collaboration to solve complex challenges, ensuring high-quality, user-focused "
            "results.",
            "Improved operational efficiency through process enhancements and technical strategy execution.",
            "Implemented observability tools, including Loki and Grafana, into the serverless platform to enable "
            "continuous monitoring, streamline issue detection, and enhance system reliability."
        ]
    },
    {
        "title": "Senior Software Engineer",
        "company": "Fidelity Investment (Contract)",
        "location": "Remote",
        "dates": "Apr 2024 – Jun 2024",
        "bullets": [
            "Developed high-performance algorithms for trading and cryptocurrency workflows, streamlining buy/sell "
            "transactions.",
            "Promoted CI/CD practices for efficient and reliable software delivery.",
            "Mentored junior engineers and aligned solutions with strategic business objectives."
        ]
    },
    {
        "title": "Backend Software Engineer",
        "company": "VMware Carbon Black",
        "location": "Remote",
        "dates": "May 2022 – Apr 2024",
        "bullets": [
            "Designed scalable microservices architecture for data pipelines, enhancing reliability and reducing "
            "latency.",
            "Led feature development for high-throughput event processing and alerting systems.",
            "Optimized pipeline performance with modern frameworks and ensured alignment with evolving security "
            "requirements."
        ]
    },
    {
        "title": "Senior Consultant",
        "company": "Lincoln Financial Group",
        "location": "Remote",
        "dates": "Dec 2019 – May 2022",
        "bullets": [
            "Transformed contact center applications into scalable microservices on AWS, improving performance and reliability.",
            "Led proof-of-concept initiatives, designed Admin UI applications, and streamlined configuration management.",
            "Conducted code reviews and mentored junior engineers to foster team development."
        ]
    },
    {
        "title": "Senior Software Engineer",
        "company": "J.B. Hunt Transport Services",
        "location": "Lowell, Arkansas",
        "dates": "Mar 2018 – Dec 2019",
        "bullets": [
            "Managed order processing for intermodal and truckload operations using Spring Boot and cloud infrastructure.",
            "Developed resilient microservices with Netflix OSS stack, improving operational efficiency.",
            "Provided mentorship and implemented high-quality, scalable solutions."
        ]
    },
    {
        "title": "Consultant",
        "company": "Syntel",
        "location": "Atlanta, Georgia",
        "dates": "Oct 2016 – Feb 2018",
        "bullets": [
            "Designed and developed responsive web applications with microservices for enhanced scalability.",
            "Contributed to system design solutions and mentored team members.",
            "Engage in Pair Programming sessions and adhere to Test-Driven Development (TDD) practices to ensure code quality, reliability, and maintainability throughout the development lifecycle."
        ]
    },
    {
        "title": "Senior Software Engineer",
        "company": "Photon Infotech",
        "location": "Philadelphia, Pennsylvania",
        "dates": "Dec 2014 – Oct 2016",
        "bullets": [
            "Designed and developed responsive web applications with microservices for enhanced scalability.",
            "Contributed to system design solutions and mentored team members.",
            "Engage in Pair Programming sessions and adhere to Test-Driven Development (TDD) practices to ensure code "
            "quality, reliability, and maintainability throughout the development lifecycle."
        ]
    },
    {
        "title": "Software Engineer",
        "company": "LTI",
        "location": "Chennai, Tamil Nadu",
        "dates": "Oct 2011 – Dec 2014",
        "bullets": [
            "Led full-stack development for scalable web applications, focusing on microservices and Agile "
            "methodologies. ",
        ]
    },
    {
        "title": "Software Developer",
        "company": "IndigoTx Software Technologies",
        "location": "Chennai, Tamil Nadu",
        "dates": "Apr 2010 – Oct 2011",
        "bullets": [
            "Led full-stack development for scalable web applications, focusing on microservices and Agile "
            "methodologies. ",
        ]
    }
]

# Add roles to the document
for role in roles:
    doc.add_heading(f"{role['title']} – {role['company']}", level=3)
    doc.add_paragraph(f"{role['location']} | {role['dates']}")
    for bullet in role["bullets"]:
        doc.add_paragraph(bullet, style="List Bullet")

# Add Education & Certifications Section
doc.add_heading("Education & Certifications", level=2)
doc.add_paragraph("Bachelor of Engineering – Computer Science, Anna University")
doc.add_paragraph("AWS Certified Solutions Architect – Associate")

# Save the document
file_path = "./Saravanakumar_Resume.docx"
doc.save(file_path)
print("Convert to PDF")

convert(file_path)

